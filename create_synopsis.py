#!/usr/bin/env python3
"""
Creates the Project Synopsis .docx matching the exact format of 'Synopsis format.docx'
with all content filled in for the NLP project.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            element = OxmlElement(f'w:{edge}')
            for key in ('sz', 'val', 'color', 'space'):
                if key in edge_data:
                    element.set(qn(f'w:{key}'), str(edge_data[key]))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_formatted_run(paragraph, text, bold=False, size=12, font_name='Times New Roman'):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.insert(0, rFonts)
    return run

def set_paragraph_spacing(paragraph, before=0, after=0, line=240):
    pPr = paragraph._element.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    spacing.set(qn('w:line'), str(line))
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

def write_cell(cell, text, bold=False, size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = alignment
    set_paragraph_spacing(p, before=0, after=0, line=240)
    add_formatted_run(p, text, bold=bold, size=size)

def write_cell_bullets(cell, items, bold=False, size=12):
    cell.text = ''
    for i, item in enumerate(items):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0, line=240)
        add_formatted_run(p, f"  {i+1}. {item}", bold=bold, size=size)

def write_cell_bullet_dots(cell, items, size=12):
    cell.text = ''
    for i, item in enumerate(items):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0, line=240)
        add_formatted_run(p, f"  \u2022 {item}", size=size)

def merge_cells_in_row(table, row_idx, start_col, end_col):
    cell_start = table.cell(row_idx, start_col)
    cell_end = table.cell(row_idx, end_col)
    cell_start.merge(cell_end)

def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ─── Header: School Name ───
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('School of Computer Science and Engineering')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    # ─── Sub-header: Project Synopsis ───
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Project Synopsis')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.underline = True

    # ─── Main Table ───
    # Columns: Label | Content-Left | Content-Right | Sign
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(6.5)
        row.cells[2].width = Cm(4.0)
        row.cells[3].width = Cm(2.5)

    # ─── ROW: Team Members Header ───
    row = table.rows[0]
    write_cell(row.cells[0], 'Project Team Members:\n\n\nTeam Id: _______', bold=True, size=12)

    write_cell(row.cells[1], 'Student Name', bold=True, size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    write_cell(row.cells[2], 'USN', bold=True, size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    write_cell(row.cells[3], 'Sign', bold=True, size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # ─── ROW: Member 1 ───
    row = table.add_row()
    write_cell(row.cells[0], '', size=12)
    write_cell(row.cells[1], '  1. __________', size=12)
    write_cell(row.cells[2], '', size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    write_cell(row.cells[3], '', size=12)

    # ─── ROW: Member 2 ───
    row = table.add_row()
    write_cell(row.cells[0], '', size=12)
    write_cell(row.cells[1], '  2. __________', size=12)
    write_cell(row.cells[2], '', size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    write_cell(row.cells[3], '', size=12)

    # ─── ROW: Member 3 ───
    row = table.add_row()
    write_cell(row.cells[0], '', size=12)
    write_cell(row.cells[1], '  3. __________', size=12)
    write_cell(row.cells[2], '', size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    write_cell(row.cells[3], '', size=12)

    # ─── ROW: Member 4 ───
    row = table.add_row()
    write_cell(row.cells[0], '', size=12)
    write_cell(row.cells[1], '  4. __________', size=12)
    write_cell(row.cells[2], '', size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    write_cell(row.cells[3], '', size=12)

    # Merge the first column (Team Members label) for rows 0-4
    table.cell(0, 0).merge(table.cell(4, 0))
    write_cell(table.cell(0, 0), 'Project Team Members:\n\n\n\nTeam Id: _______', bold=True, size=12)

    # ─── ROW: Guide Name ───
    row = table.add_row()
    write_cell(row.cells[0], 'Guide Name', bold=True, size=12)
    # Merge cols 1-2 for guide name
    row.cells[1].merge(row.cells[2])
    write_cell(row.cells[1], 'Dr./Prof. __________', size=12)
    write_cell(row.cells[3], '', size=12)

    # ─── ROW: Title ───
    row = table.add_row()
    write_cell(row.cells[0], 'Title of the Project', bold=True, size=12)
    row.cells[1].merge(row.cells[3])
    write_cell(row.cells[1],
        'Multilingual Content Integrity and Authorship Intelligence Platform: '
        'An Ensemble Transformer Approach for AI-Generated Text Detection, '
        'Semantic Plagiarism Identification, and Adaptive Content Humanization',
        bold=True, size=12)

    # ─── ROW: Objective/Purpose ───
    row = table.add_row()
    write_cell(row.cells[0], 'Objective / Purpose', bold=True, size=12)
    row.cells[1].merge(row.cells[3])
    cell = row.cells[1]
    cell.text = ''
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, line=240)
    add_formatted_run(p, 'The objectives of the proposed work are as follows:', size=12)

    objectives = [
        'To develop a high-accuracy multilingual AI-generated text detection system using an ensemble of fine-tuned transformer models (DeBERTa-v3, RoBERTa, Longformer, XLM-RoBERTa) capable of identifying machine-generated content across 100+ languages with a target accuracy exceeding 95%.',
        'To build a semantic plagiarism detection engine that goes beyond traditional n-gram matching by employing sentence-level embeddings (Sentence-BERT, SimCSE) and cross-encoder verification (DeBERTa) to detect paraphrase-level and meaning-level plagiarism.',
        'To design an adaptive content humanization module using sequence-to-sequence models (Flan-T5-XL, PEGASUS, DIPPER, Mistral-7B with QLoRA) that transforms AI-generated text into naturally human-written content, achieving near-zero AI detection scores while preserving semantic meaning.',
        'To integrate all three modules into a unified, self-contained platform that operates entirely offline without dependency on third-party APIs, providing transparent and reproducible content analysis through a command-line interface and web application.',
    ]
    for i, obj in enumerate(objectives):
        p = cell.add_paragraph()
        set_paragraph_spacing(p, line=240)
        add_formatted_run(p, f'  {i+1}. {obj}', size=12)

    # ─── ROW: Abstract ───
    row = table.add_row()
    cell_label = row.cells[0]
    cell_label.text = ''
    p = cell_label.paragraphs[0]
    add_formatted_run(p, 'Abstract', bold=True, size=12)
    p2 = cell_label.add_paragraph()
    add_formatted_run(p2, '(Approximately 10 lines)', size=9)

    row.cells[1].merge(row.cells[3])
    write_cell(row.cells[1],
        'The rapid proliferation of large language models has created an urgent need for reliable systems '
        'that can detect AI-generated content, identify plagiarism, and ensure content authenticity across '
        'multiple languages. Existing commercial tools suffer from high false-positive rates, limited '
        'multilingual support, and inability to detect paraphrase-level plagiarism. This project proposes '
        'a unified multilingual content intelligence platform built on an ensemble of fine-tuned transformer '
        'architectures. The AI detection module employs four specialized classifiers \u2014 DeBERTa-v3-large, '
        'RoBERTa-large, Longformer-base for long documents, and XLM-RoBERTa-large for multilingual text '
        '\u2014 combined through a meta-classifier to maximize detection accuracy across diverse text formats '
        'and languages. The plagiarism detection engine uses MinHash/LSH for fast candidate retrieval '
        'followed by Sentence-BERT and SimCSE embeddings for semantic similarity scoring, with a DeBERTa '
        'cross-encoder for verification. The humanization module fine-tunes DIPPER, Flan-T5-XL, PEGASUS, '
        'and Mistral-7B to rewrite flagged content into naturally human-like text, employing an iterative '
        'feedback loop with the detection module to ensure transformed output registers minimal AI signatures. '
        'The system is trained on 18 large-scale datasets including RAID (10M+ documents), M4 (multilingual), '
        'HC3, and PAN corpora, and operates entirely locally without external API dependencies.',
        size=12)

    # ─── ROW: Hardware & Software ───
    row = table.add_row()
    write_cell(row.cells[0], 'Hardware & Software Requirements', bold=True, size=12)

    # Merge cols 1 for hardware
    cell_hw = row.cells[1]
    cell_hw.text = ''
    p = cell_hw.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_run(p, 'Hardware', size=12)
    run = p.runs[0]
    run.underline = True

    hw_items = [
        'GPU: NVIDIA RTX 4090 (24 GB) or A100 (40/80 GB)',
        'RAM: 64 GB minimum',
        'Storage: 500 GB+ NVMe SSD',
        'CPU: 8+ core (Intel i7/i9 or AMD Ryzen)',
        'Cloud: RunPod / Lambda Labs / Colab Pro+',
    ]
    for item in hw_items:
        p = cell_hw.add_paragraph()
        set_paragraph_spacing(p, line=240)
        add_formatted_run(p, f'  \u2022 {item}', size=12)

    # Merge cols 2-3 for software
    row.cells[2].merge(row.cells[3])
    cell_sw = row.cells[2]
    cell_sw.text = ''
    p = cell_sw.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_run(p, 'Software', size=12)
    run = p.runs[0]
    run.underline = True

    sw_items = [
        'Python 3.10+',
        'PyTorch 2.0+ with CUDA 12.1+',
        'Hugging Face Transformers 4.36+',
        'Sentence-Transformers library',
        'Datasketch (MinHash/LSH)',
        'scikit-learn',
        'NLTK / spaCy',
        'FastAPI + HTML/CSS/JS',
        'Weights & Biases / TensorBoard',
    ]
    for item in sw_items:
        p = cell_sw.add_paragraph()
        set_paragraph_spacing(p, line=240)
        add_formatted_run(p, f'  \u2022 {item}', size=12)

    # ─── ROW: Innovation/Novelty ───
    row = table.add_row()
    write_cell(row.cells[0], 'Innovation / Novelty Involved', bold=True, size=12)
    row.cells[1].merge(row.cells[3])
    cell = row.cells[1]
    cell.text = ''

    innovations = [
        'Unified Three-Module Architecture: Integrates AI detection, plagiarism checking, and content humanization into a single pipeline sharing a common semantic understanding layer, unlike existing tools that treat these as separate problems.',
        'Multilingual Ensemble Detection: Employs XLM-RoBERTa-large trained on the M4 multilingual dataset, enabling AI detection across 100+ languages \u2014 a capability absent in most commercial detectors.',
        'Feedback-Loop Humanization: Uses an iterative refinement loop where humanized output is continuously evaluated by the detection module, adjusting diversity and reordering parameters until AI detection scores fall below the target threshold.',
        'Semantic Plagiarism Detection: Combines MinHash/LSH fast screening with Sentence-BERT/SimCSE embeddings and DeBERTa cross-encoder verification to detect meaning-level plagiarism that string-matching tools miss.',
        'Fully Self-Contained and Transparent: All models trained locally on open datasets with no proprietary API dependency, making the system reproducible, auditable, and offline-capable.',
    ]
    for i, item in enumerate(innovations):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        set_paragraph_spacing(p, line=240)
        add_formatted_run(p, f'  {i+1}. {item}', size=12)

    # ─── ROW: Algorithms/Technology ───
    row = table.add_row()
    write_cell(row.cells[0], 'Algorithms / Technology Planned', bold=True, size=12)
    row.cells[1].merge(row.cells[3])
    cell = row.cells[1]
    cell.text = ''

    algorithms = [
        'Transformer-based Sequence Classification (DeBERTa-v3, RoBERTa, Longformer, XLM-RoBERTa)',
        'Ensemble Meta-Classification using Logistic Regression',
        'MinHash / Locality-Sensitive Hashing (LSH) for fast plagiarism screening',
        'Contrastive Learning for Sentence Embeddings (Sentence-BERT, SimCSE)',
        'Cross-Encoder Pairwise Verification (DeBERTa-v3)',
        'Sequence-to-Sequence Text Generation (Flan-T5-XL, PEGASUS-large)',
        'Parameter-Efficient Fine-Tuning with QLoRA (Mistral-7B)',
        'Discourse-Level Paraphrasing (DIPPER) with controllable diversity',
        'Iterative Feedback Refinement Loop for humanization quality assurance',
    ]
    for i, item in enumerate(algorithms):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        set_paragraph_spacing(p, line=240)
        add_formatted_run(p, f'  {i+1}. {item}', size=12)

    # ─── ROW: Methodology ───
    row = table.add_row()
    cell_label = row.cells[0]
    cell_label.text = ''
    p = cell_label.paragraphs[0]
    add_formatted_run(p, 'Methodology', bold=True, size=12)
    p2 = cell_label.add_paragraph()
    add_formatted_run(p2, '(in bullets)', size=10)

    row.cells[1].merge(row.cells[3])
    cell = row.cells[1]
    cell.text = ''

    methodology = [
        'Collect and preprocess 18 large-scale NLP datasets spanning AI detection (RAID, HC3, M4, GPT-2 Output, FAIDSet, PAN Author ID), plagiarism detection (PAN Plagiarism Corpora, Clough & Stevenson, Webis, WikiSplit, STS Benchmark, PAWS), and humanization (ParaNMT-50M, QQP, MRPC, BEA-2019 GEC).',
        'Convert all datasets into unified formats: classification (text + label), paraphrase (input + output), and similarity (text_a + text_b + score). Apply text cleaning, deduplication, and stratified 80/10/10 train/validation/test splits.',
        'Fine-tune four transformer classifiers for AI detection: DeBERTa-v3-large as primary detector, RoBERTa-large as ensemble member, Longformer-base for long documents (up to 4096 tokens), and XLM-RoBERTa-large for multilingual content.',
        'Train a logistic regression meta-classifier on the combined prediction outputs of all four models to produce an optimized ensemble AI detection score.',
        'Implement MinHash/LSH-based reference indexing for fast plagiarism candidate retrieval, followed by Sentence-BERT and SimCSE sentence-level embedding comparison.',
        'Fine-tune a DeBERTa-v3 cross-encoder for precise pairwise plagiarism verification on flagged sentence pairs.',
        'Fine-tune Flan-T5-XL, PEGASUS-large, and Mistral-7B (via QLoRA) on paraphrase datasets for content humanization. Configure DIPPER for paragraph-level paraphrasing with controllable diversity.',
        'Implement an iterative feedback loop: pass humanized output through the AI detector, and if the score exceeds the threshold, increase diversity/reordering parameters and regenerate until the target is met.',
        'Integrate all three modules into a unified pipeline with a CLI interface supporting --detect, --plagiarism, --humanize, and --full flags.',
        'Build a web application using FastAPI (backend) and HTML/CSS/JS (frontend) for user-friendly text analysis and transformation.',
        'Evaluate the complete system using accuracy, F1-score, precision, recall, and AUROC metrics on held-out test sets.',
    ]
    for i, item in enumerate(methodology):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        set_paragraph_spacing(p, line=240)
        add_formatted_run(p, f'  \u2022 {item}', size=12)

    # ─── ROW: Expected Outcomes ───
    row = table.add_row()
    write_cell(row.cells[0], 'Expected Outcomes', bold=True, size=12)
    row.cells[1].merge(row.cells[3])
    cell = row.cells[1]
    cell.text = ''
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, line=240)
    add_formatted_run(p, 'The expected outcomes are as follows:', size=12)

    outcomes = [
        'An AI detection ensemble achieving 95%+ accuracy on standard benchmarks (RAID, HC3, M4) with robust performance across multiple languages, text lengths, and AI generators, while maintaining a false positive rate below 5%.',
        'A semantic plagiarism detection system capable of identifying copy-paste, paraphrase-level, and meaning-level plagiarism, evaluated on PAN benchmark corpora with competitive precision and recall scores.',
        'A content humanization module that transforms AI-generated text to achieve near-zero (< 5%) AI detection scores on both the internal detector and external commercial tools, while preserving over 85% semantic similarity with the original content.',
        'A fully integrated, self-contained platform with CLI and web interfaces that processes text through all three modules in under 30 seconds per document, operating entirely offline without external API dependencies.',
    ]
    for i, item in enumerate(outcomes):
        p = cell.add_paragraph()
        set_paragraph_spacing(p, line=240)
        add_formatted_run(p, f'  {i+1}. {item}', size=12)

    # ─── ROW: Literature Survey (empty label for extra space) ───
    row = table.add_row()
    write_cell(row.cells[0], 'Literature Survey\n(Minimum 10 Papers)', bold=True, size=12)
    row.cells[1].merge(row.cells[3])
    cell = row.cells[1]
    cell.text = ''

    papers = [
        '[1] Dugan, L., et al. "RAID: A Shared Benchmark for Robust Evaluation of Machine-Generated Text Detectors." ACL 2024. arXiv:2405.07940',
        '[2] Wang, Y., et al. "M4: Multi-generator, Multi-domain, Multi-lingual Black-Box Machine-Generated Text Detection." 2024. arXiv:2305.14902',
        '[3] Guo, B., et al. "HC3: Human ChatGPT Comparison Corpus." 2023. Hello-SimpleAI/HC3, Hugging Face.',
        '[4] Sadasivan, V.S., et al. "On the Reliability of AI-Text Detectors." 2023. arXiv:2304.02819',
        '[5] Tang, R., et al. "Detecting Machine-Generated Text: A Critical Survey." 2023. arXiv:2303.07205',
        '[6] Reimers, N. & Gurevych, I. "Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks." EMNLP 2019. arXiv:1908.10084',
        '[7] Gao, T., Yao, X., & Chen, D. "SimCSE: Simple Contrastive Learning of Sentence Embeddings." EMNLP 2021. arXiv:2104.08821',
        '[8] Krishna, K., et al. "Paraphrasing Evades Detectors of AI-Generated Text, but Retrieval is an Effective Defense." (DIPPER) 2023. arXiv:2303.13408',
        '[9] He, P., et al. "DeBERTa: Decoding-enhanced BERT with Disentangled Attention." ICLR 2021. arXiv:2006.03654',
        '[10] Hlavnova, E. & Pikuliak, M. "Fine-tuned LLMs for Multilingual Machine-Generated Text Detection." SemEval-2024 Task 8. arXiv:2402.13671',
        '[11] "Fine-grained AI-generated Text Detection using Multi-task Auxiliary and Multi-level Contrastive Learning." (FAIDSet) 2025. arXiv:2505.14271',
        '[12] Foltynek, T., et al. "A Survey on Plagiarism Detection." ACM Computing Surveys, 2019. arXiv:1703.05546',
        '[13] "Authorship Style Transfer with Policy Optimization." 2024. arXiv:2403.08043',
    ]
    for i, paper in enumerate(papers):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        set_paragraph_spacing(p, line=276)
        add_formatted_run(p, paper, size=11)

    # ─── Signature section ───
    doc.add_paragraph('')
    doc.add_paragraph('')

    p = doc.add_paragraph()
    run1 = p.add_run('Signature of the Guide')
    run1.font.size = Pt(12)
    run1.font.name = 'Times New Roman'
    run2 = p.add_run('\t\t\t\t\t\tProject Coordinator')
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'

    # Save
    output_path = 'Project_Synopsis_Filled.docx'
    doc.save(output_path)
    print(f'Synopsis saved to: {output_path}')

if __name__ == '__main__':
    main()
