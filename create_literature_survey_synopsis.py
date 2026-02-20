#!/usr/bin/env python3
"""
Creates Literature Survey in RV University format:
  1. Paperwise comparison table (Title, Author, Year, Objective/Purpose, Methodology/Technique, Findings)
  2. Each research paper elaborated in its own paragraph with arXiv hyperlink
Output: Literature_Survey_Synopsis.docx
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml


def fmt_run(paragraph, text, bold=False, italic=False, size=12, font='Times New Roman', color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def add_hyperlink(paragraph, text, url, size=12, font='Times New Roman'):
    """Add a clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(size * 2))
    rPr.append(sz)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)
    rPr.append(rFonts)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)


def write_cell(cell, text, bold=False, size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = alignment
    fmt_run(p, text, bold=bold, size=size)


# ══════════════════════════════════════════════════════════════
# Paper data: all 13 references
# ══════════════════════════════════════════════════════════════
PAPERS = [
    {
        'title': 'RAID: A Shared Benchmark for Robust Evaluation of Machine-Generated Text Detectors',
        'authors': 'Dugan, L., Hwang, A., Trhlik, F., et al.',
        'year': '2024',
        'objective': 'To create a large-scale shared benchmark for evaluating AI text detectors across multiple generators, domains, and adversarial attacks.',
        'methodology': 'Compiled 6M+ generated documents from 11 language models across 8 domains with 11 adversarial attack strategies. Evaluated existing detectors on this unified benchmark.',
        'findings': 'Most existing detectors showed significant performance drops under adversarial attacks. The benchmark revealed that no single detector generalizes well across all generators and domains, highlighting the need for ensemble approaches.',
        'url': 'https://arxiv.org/abs/2405.07940',
        'paragraph': 'Dugan et al. [1] introduced RAID, the largest shared benchmark for evaluating machine-generated text detectors. The dataset contains over six million generated documents spanning eleven different language models, eight subject domains, and eleven adversarial attack strategies. RAID addresses a critical gap in the field where detectors were previously evaluated on small, inconsistent datasets, making cross-study comparison unreliable. Their evaluation of existing detectors on this benchmark revealed that most systems suffer significant accuracy drops when faced with adversarial attacks such as paraphrasing and homoglyph substitution. The work was published at ACL 2024 and has become a standard reference point for benchmarking new detection systems. For our project, RAID serves as the primary evaluation benchmark for the AI detection ensemble.',
    },
    {
        'title': 'M4: Multi-generator, Multi-domain, Multi-lingual Black-Box Machine-Generated Text Detection',
        'authors': 'Wang, Y., et al. (MBZUAI)',
        'year': '2024',
        'objective': 'To build a multilingual, multi-generator dataset and detection framework for cross-lingual AI text detection.',
        'methodology': 'Collected machine-generated text from multiple LLMs across multiple languages and domains. Trained and evaluated detection models in both monolingual and cross-lingual settings.',
        'findings': 'Detectors trained exclusively on English data perform poorly on other languages. Cross-lingual transfer learning with multilingual transformers like XLM-RoBERTa significantly improves detection across languages.',
        'url': 'https://arxiv.org/abs/2305.14902',
        'paragraph': 'Wang et al. [2] presented M4, a multilingual and multi-generator dataset that addresses the reality that AI-generated content is not limited to English. The dataset covers text produced by multiple large language models across several languages and domains. Their experiments demonstrated that detectors trained solely on English data perform poorly when applied to other languages, a finding that has significant implications for global deployment of detection tools. The cross-lingual evaluation showed that multilingual transformer models like XLM-RoBERTa can bridge this gap when fine-tuned on diverse multilingual data. This work directly motivated our decision to include XLM-RoBERTa-large in the detection ensemble, trained on the M4 corpus to ensure robust multilingual coverage.',
    },
    {
        'title': 'HC3: Human ChatGPT Comparison Corpus',
        'authors': 'Guo, B., et al. (Hello-SimpleAI)',
        'year': '2023',
        'objective': 'To create a paired dataset of human-written and ChatGPT-generated answers for training and evaluating AI text detectors.',
        'methodology': 'Collected question-answer pairs from multiple domains where both human experts and ChatGPT provided answers. Analyzed linguistic differences between human and AI responses.',
        'findings': 'ChatGPT responses tend to be more verbose, use more formal language, and exhibit less variability than human answers. Simple classifiers trained on HC3 achieve high accuracy but may not generalize to other generators.',
        'url': 'https://huggingface.co/datasets/Hello-SimpleAI/HC3',
        'paragraph': 'Guo et al. [3] assembled HC3, the Human ChatGPT Comparison Corpus, which pairs human-written answers with ChatGPT responses across several question-answering domains including finance, medicine, and open-domain QA. The paired structure makes HC3 particularly useful for training binary classifiers because it provides direct contrastive examples of human versus AI writing on the same topics. Their analysis revealed that ChatGPT responses tend to be more structured, verbose, and formally written compared to human answers. However, since HC3 is limited to a single generator (ChatGPT), models trained exclusively on it may not generalize to text from other AI systems. In our project, HC3 serves as a complementary training dataset alongside the broader RAID and M4 benchmarks.',
    },
    {
        'title': 'On the Reliability of AI-Text Detectors',
        'authors': 'Sadasivan, V.S., et al.',
        'year': '2023',
        'objective': 'To investigate the fundamental limitations of AI text detection methods and assess their robustness against evasion attacks.',
        'methodology': 'Tested multiple detection methods (watermarking, neural classifiers, statistical methods) against paraphrasing attacks using models like DIPPER. Provided theoretical analysis of detection limits.',
        'findings': 'Simple paraphrasing attacks can reduce detection accuracy to near-random levels. As language models improve, the statistical gap between human and machine text shrinks, making detection inherently harder.',
        'url': 'https://arxiv.org/abs/2304.02819',
        'paragraph': 'Sadasivan et al. [4] raised fundamental questions about the reliability of AI text detectors. Their experiments demonstrated that simple paraphrasing attacks can dramatically reduce detection accuracy, sometimes bringing it close to random chance. They tested multiple detection paradigms including watermarking, neural classifiers, and statistical methods, finding that none were robust against determined adversaries. Their theoretical analysis argued that as language models continue to improve, the distributional gap between human and machine text will shrink further, making detection an increasingly difficult problem. This work served as a wake-up call for the research community and directly influenced our decision to include adversarial robustness testing in the evaluation pipeline and to use an ensemble of diverse architectures rather than relying on any single detector.',
    },
    {
        'title': 'Detecting Machine-Generated Text: A Critical Survey',
        'authors': 'Tang, R., et al.',
        'year': '2023',
        'objective': 'To provide a comprehensive survey of machine-generated text detection methods, categorizing existing approaches and identifying open challenges.',
        'methodology': 'Systematic review of detection methods categorized into statistical approaches, neural classifiers, and watermarking-based techniques. Comparative analysis across different evaluation settings.',
        'findings': 'No single detection method dominates across all settings. Statistical methods work for older generators but struggle with newer fluent models. Neural classifiers achieve higher accuracy but are brittle under distribution shift. Ensemble methods offer the best generalization.',
        'url': 'https://arxiv.org/abs/2303.07205',
        'paragraph': 'Tang et al. [5] provided a thorough survey of machine-generated text detection methods, organizing them into three main categories: statistical approaches based on perplexity and entropy, neural classifiers fine-tuned on labelled data, and watermarking-based techniques that embed detectable signals during generation. Their comparative analysis highlighted that no single method dominates across all evaluation settings. Statistical methods like perplexity scoring work well for older generators like GPT-2 but struggle with newer models that produce more fluent text. Neural classifiers achieve higher accuracy but can be brittle when the test distribution shifts away from training data. This finding reinforced our architectural choice to use an ensemble of four different transformer models rather than relying on any single approach.',
    },
    {
        'title': 'Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks',
        'authors': 'Reimers, N. & Gurevych, I.',
        'year': '2019',
        'objective': 'To derive semantically meaningful sentence embeddings from BERT that can be compared using cosine similarity for efficient similarity search.',
        'methodology': 'Modified BERT with Siamese and triplet network structures. Trained on NLI and STS datasets using contrastive and triplet loss functions to produce fixed-size sentence embeddings.',
        'findings': 'Sentence-BERT reduces the computational cost of finding similar sentence pairs from quadratic (65 hours for 10K sentences with BERT cross-encoder) to linear (5 seconds with cosine similarity), while maintaining competitive accuracy on STS benchmarks.',
        'url': 'https://arxiv.org/abs/1908.10084',
        'paragraph': 'Reimers and Gurevych [6] proposed Sentence-BERT (SBERT), which modifies the standard BERT architecture with Siamese and triplet network structures to produce fixed-size sentence embeddings that can be compared using cosine similarity. This was a breakthrough for practical applications because it reduced the computational cost of pairwise sentence comparison from quadratic to linear complexity. For example, finding the most similar pair among 10,000 sentences takes 65 hours with a BERT cross-encoder but only 5 seconds with SBERT embeddings. The model was trained on Natural Language Inference (NLI) and Semantic Textual Similarity (STS) datasets. In our plagiarism detection module, Sentence-BERT serves as the primary embedding model for computing semantic similarity between query documents and the reference corpus.',
    },
    {
        'title': 'SimCSE: Simple Contrastive Learning of Sentence Embeddings',
        'authors': 'Gao, T., Yao, X., & Chen, D.',
        'year': '2021',
        'objective': 'To improve sentence embedding quality through a simple contrastive learning framework using both unsupervised and supervised approaches.',
        'methodology': 'Unsupervised: uses dropout as minimal data augmentation, passing the same sentence through the encoder twice with different dropout masks. Supervised: leverages NLI pairs with contrastive objectives.',
        'findings': 'SimCSE achieves significantly better alignment and uniformity in the embedding space compared to previous methods. The supervised variant outperforms Sentence-BERT on multiple STS benchmarks.',
        'url': 'https://arxiv.org/abs/2104.08821',
        'paragraph': 'Gao, Yao, and Chen [7] introduced SimCSE, a contrastive learning framework that significantly improves sentence embedding quality. The unsupervised variant uses an elegantly simple approach: passing the same sentence through the encoder twice with different dropout masks to create positive pairs. The supervised variant leverages Natural Language Inference pairs, pushing entailment examples closer together and contradiction examples further apart in the embedding space. SimCSE achieves better alignment and uniformity properties compared to previous methods including Sentence-BERT. We use SimCSE-enhanced embeddings alongside Sentence-BERT in our plagiarism detection pipeline to improve recall, particularly for cases where the wording has been substantially changed but the underlying meaning is preserved.',
    },
    {
        'title': 'Paraphrasing Evades Detectors of AI-Generated Text, but Retrieval is an Effective Defense (DIPPER)',
        'authors': 'Krishna, K., et al.',
        'year': '2023',
        'objective': 'To demonstrate that discourse-level paraphrasing can evade AI text detectors and to propose retrieval-based detection as a robust countermeasure.',
        'methodology': 'Developed DIPPER, an 11B parameter T5-XXL paraphraser with controllable lexical diversity and order diversity knobs. Tested evasion against multiple detectors and proposed retrieval-based defense.',
        'findings': 'Even moderate paraphrasing with DIPPER fools most existing detectors. However, retrieval-based methods that compare against a database of known AI outputs remain effective even after paraphrasing.',
        'url': 'https://arxiv.org/abs/2303.13408',
        'paragraph': 'Krishna et al. [8] introduced DIPPER (Discourse Paraphrase Rewriter), an 11-billion parameter T5-XXL model fine-tuned for paragraph-level paraphrasing with two controllable parameters: lexical diversity (how much vocabulary changes) and order diversity (how much sentence structure is rearranged). Their key finding was that even moderate paraphrasing could fool most existing AI text detectors, but retrieval-based methods that compare text against a database of known AI outputs remained effective. This dual insight is central to our project: DIPPER serves as one of the rewriting models in the humanization module for transforming AI-generated text, while the retrieval finding informed our decision to combine the detection module with semantic similarity search in the plagiarism detection pipeline.',
    },
    {
        'title': 'DeBERTa: Decoding-enhanced BERT with Disentangled Attention',
        'authors': 'He, P., Liu, X., Gao, J., & Chen, W.',
        'year': '2021',
        'objective': 'To improve upon BERT by introducing disentangled attention that separately encodes content and position information, along with an enhanced mask decoder.',
        'methodology': 'Proposed two novel techniques: disentangled attention mechanism using separate vectors for content and position, and an enhanced mask decoder that incorporates absolute position information for token prediction.',
        'findings': 'DeBERTa achieved state-of-the-art results on multiple NLU benchmarks including SuperGLUE. The disentangled attention provides better fine-grained token-level understanding compared to standard BERT attention.',
        'url': 'https://arxiv.org/abs/2006.03654',
        'paragraph': 'He et al. [9] proposed DeBERTa (Decoding-enhanced BERT with Disentangled Attention), which introduces two key innovations over the original BERT architecture. First, the disentangled attention mechanism uses separate vectors to encode content and position information, allowing the model to better capture the relationship between word meaning and position. Second, an enhanced mask decoder incorporates absolute position information in the decoding layer for more accurate token prediction. DeBERTa achieved state-of-the-art results on multiple natural language understanding benchmarks including SuperGLUE. In our project, DeBERTa-v3-large serves as the primary classifier in the AI detection ensemble due to its superior fine-grained understanding of text patterns that distinguish human from machine writing.',
    },
    {
        'title': 'Fine-tuned LLMs for Multilingual Machine-Generated Text Detection (SemEval-2024 Task 8)',
        'authors': 'Hlavnova, E. & Pikuliak, M.',
        'year': '2024',
        'objective': 'To explore fine-tuning strategies for detecting machine-generated text across multiple languages using parameter-efficient methods.',
        'methodology': 'Fine-tuned RoBERTa and XLM-RoBERTa models using LoRA (Low-Rank Adaptation) for multilingual AI text detection. Used majority voting across language-specific classification heads.',
        'findings': 'LoRA-adapted RoBERTa with majority voting achieves competitive multilingual detection performance without training separate models per language. Parameter-efficient fine-tuning is practical for multilingual deployment.',
        'url': 'https://arxiv.org/abs/2402.13671',
        'paragraph': 'Hlavnova and Pikuliak [10] explored fine-tuning strategies for multilingual machine-generated text detection as part of SemEval-2024 Task 8. They found that LoRA-adapted (Low-Rank Adaptation) RoBERTa models combined with majority voting across language-specific classification heads could achieve competitive performance without the need to train separate models for each language. Their results confirmed that parameter-efficient fine-tuning is a practical and scalable path for multilingual detection, which aligns with our use of XLM-RoBERTa-large trained on the M4 multilingual corpus. The majority voting strategy they proposed also informed our meta-classifier design in the detection ensemble.',
    },
    {
        'title': 'Fine-grained AI-generated Text Detection using Multi-task Auxiliary and Multi-level Contrastive Learning (FAIDSet)',
        'authors': 'Multiple authors',
        'year': '2025',
        'objective': 'To introduce a fine-grained detection framework that classifies text as fully human, fully AI, or mixed content, with a large-scale multilingual dataset.',
        'methodology': 'Multi-task learning with auxiliary sentence-level detection alongside document-level classification. Multi-level contrastive learning to separate human, AI, and mixed text representations.',
        'findings': 'The fine-grained three-class approach (human/AI/mixed) outperforms binary classification on real-world data where documents often contain both human and AI-written sections. The multilingual dataset enables cross-lingual fine-grained detection.',
        'url': 'https://arxiv.org/abs/2505.14271',
        'paragraph': 'The FAIDSet framework [11] proposed a fine-grained labelling scheme that goes beyond the traditional binary human-or-AI distinction. Instead of treating detection as a two-class problem, FAIDSet introduces a third category for mixed content where parts of a document are human-written and parts are machine-generated. This reflects how AI tools are actually used in practice, since many writers use language models to draft certain sections while writing others themselves. The accompanying multi-task learning approach combines document-level classification with sentence-level auxiliary detection, while multi-level contrastive learning separates the three categories in the embedding space. The large-scale multilingual dataset provided with FAIDSet is used in our project for training the fine-grained detection capability.',
    },
    {
        'title': 'A Survey on Plagiarism Detection',
        'authors': 'Foltýnek, T., Meuschke, N., & Gipp, B.',
        'year': '2019',
        'objective': 'To provide a comprehensive survey of plagiarism detection methods covering string matching, citation analysis, and semantic approaches.',
        'methodology': 'Systematic review categorizing plagiarism detection into string-based, syntax-based, semantic-based, and citation-based methods. Analysis of commercial and academic detection tools.',
        'findings': 'String-based methods detect verbatim copying effectively but miss paraphrased plagiarism. Semantic approaches using embeddings are the most promising for detecting meaning-level plagiarism but were the least mature at the time of the survey.',
        'url': 'https://arxiv.org/abs/1703.05546',
        'paragraph': 'Foltýnek, Meuschke, and Gipp [12] provided a comprehensive survey of plagiarism detection methods, categorizing them into string-based, syntax-based, semantic-based, and citation-based approaches. Their analysis found that string-based methods like n-gram matching and fingerprinting are effective for detecting verbatim copying but completely miss paraphrased content. Syntax-based methods offer marginal improvement but still struggle with meaning-preserving rewrites. Semantic approaches using word and sentence embeddings were identified as the most promising direction but were the least mature at the time of writing. Our project builds directly on this observation by placing sentence-level semantic embeddings (Sentence-BERT and SimCSE) at the core of the plagiarism detection module, combined with MinHash/LSH for fast candidate retrieval.',
    },
    {
        'title': 'Authorship Style Transfer with Policy Optimization',
        'authors': 'Multiple authors',
        'year': '2024',
        'objective': 'To develop policy optimization techniques for transferring authorship style while preserving content meaning.',
        'methodology': 'Used reinforcement learning with style classifiers as reward signals to guide a language model in adopting target writing styles. Evaluated on style transfer accuracy and content preservation.',
        'findings': 'Policy optimization enables controlled style transfer that preserves semantic content while successfully changing authorship characteristics. The approach outperforms supervised fine-tuning on style transfer benchmarks.',
        'url': 'https://arxiv.org/abs/2403.08043',
        'paragraph': 'This work [13] proposed policy optimization techniques for authorship style transfer, using reinforcement learning signals from style classifiers to guide a language model in adopting a target writing style while preserving the underlying content. The approach demonstrated that RL-based optimization outperforms supervised fine-tuning for controlled style transfer tasks. While our project does not directly implement authorship style transfer, the iterative feedback loop in our humanization module draws on a similar principle: the AI detection score from our ensemble acts as a reward signal that guides the rewriting process toward more human-like output, iteratively adjusting parameters until the detection score falls below the target threshold.',
    },
]


def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ─── Header: RV University ───
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt_run(p, 'RV UNIVERSITY', bold=True, size=14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt_run(p, 'School of Computer Science and Engineering', bold=False, size=11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt_run(p, 'RV Vidyaniketan, 8th Mile, Mysuru Road, Bengaluru, 560059, India', size=10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt_run(p, 'Ph : +91 80681 99900  |  www.rvu.edu.in', size=10)

    doc.add_paragraph()

    # ─── Title ───
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fmt_run(p, 'Literature Survey Report ( Minimum 10 Research Papers )', bold=True, size=16)
    run.underline = True

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════
    # SECTION 1: Paperwise Comparison Table
    # ══════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt_run(p, '1. Paperwise Comparison Table', bold=True, size=14)

    doc.add_paragraph()

    # Create table: Title | Author | Year | Objective/Purpose | Methodology/Technique | Findings
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    headers = ['Title', 'Author', 'Year', 'Objective / Purpose', 'Methodology / Technique', 'Findings']
    hdr_row = table.rows[0]
    for i, hdr in enumerate(headers):
        write_cell(hdr_row.cells[i], hdr, bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # Data rows
    for paper in PAPERS:
        row = table.add_row()
        vals = [paper['title'], paper['authors'], paper['year'],
                paper['objective'], paper['methodology'], paper['findings']]
        for i, val in enumerate(vals):
            write_cell(row.cells[i], val, size=9)

    # Set column widths
    col_widths = [Cm(3.0), Cm(2.5), Cm(1.2), Cm(3.8), Cm(3.8), Cm(4.0)]
    for row in table.rows:
        for i, width in enumerate(col_widths):
            row.cells[i].width = width

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # SECTION 2: Elaborate Each Research Paper in Paragraph
    # ══════════════════════════════════════════════════════════════

    # Repeat header on new page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt_run(p, 'RV UNIVERSITY', bold=True, size=14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt_run(p, 'School of Computer Science and Engineering', size=11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt_run(p, 'RV Vidyaniketan, 8th Mile, Mysuru Road, Bengaluru, 560059, India', size=10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt_run(p, 'Ph : +91 80681 99900  |  www.rvu.edu.in', size=10)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt_run(p, '2. Elaborate Each Research Paper in Paragraph', bold=True, size=14)

    p = doc.add_paragraph()
    fmt_run(p, '   (1 Research Paper = 1 Paragraph, with hyperlink to the paper)', bold=True, size=11)

    doc.add_paragraph()

    # Each paper as a numbered paragraph with hyperlink
    for idx, paper in enumerate(PAPERS, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Paper number and title in bold
        fmt_run(p, f'[{idx}] ', bold=True, size=12)
        fmt_run(p, paper['title'], bold=True, italic=True, size=12)
        fmt_run(p, f' ({paper["authors"]}, {paper["year"]})', bold=False, size=12)
        fmt_run(p, '\n', size=12)

        # Paragraph elaboration
        fmt_run(p, paper['paragraph'], size=12)

        # Hyperlink on new line
        fmt_run(p, '\n\nPaper Link: ', bold=True, size=11)
        add_hyperlink(p, paper['url'], paper['url'], size=11)

        # Spacer
        doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════
    # SECTION 3: References
    # ══════════════════════════════════════════════════════════════
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt_run(p, 'References', bold=True, size=14)
    doc.add_paragraph()

    references = [
        '[1] Dugan, L., Hwang, A., Trhlik, F., et al., "RAID: A Shared Benchmark for Robust Evaluation of Machine-Generated Text Detectors," Proceedings of the Association for Computational Linguistics (ACL), 2024. arXiv:2405.07940.',
        '[2] Wang, Y., et al., "M4: Multi-generator, Multi-domain, Multi-lingual Black-Box Machine-Generated Text Detection," 2024. arXiv:2305.14902.',
        '[3] Guo, B., et al., "HC3: Human ChatGPT Comparison Corpus," Hello-SimpleAI, Hugging Face, 2023.',
        '[4] Sadasivan, V.S., et al., "On the Reliability of AI-Text Detectors," 2023. arXiv:2304.02819.',
        '[5] Tang, R., et al., "Detecting Machine-Generated Text: A Critical Survey," 2023. arXiv:2303.07205.',
        '[6] Reimers, N. and Gurevych, I., "Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks," Proceedings of EMNLP, 2019. arXiv:1908.10084.',
        '[7] Gao, T., Yao, X., and Chen, D., "SimCSE: Simple Contrastive Learning of Sentence Embeddings," Proceedings of EMNLP, 2021. arXiv:2104.08821.',
        '[8] Krishna, K., et al., "Paraphrasing Evades Detectors of AI-Generated Text, but Retrieval is an Effective Defense," (DIPPER), 2023. arXiv:2303.13408.',
        '[9] He, P., Liu, X., Gao, J., and Chen, W., "DeBERTa: Decoding-enhanced BERT with Disentangled Attention," Proceedings of ICLR, 2021. arXiv:2006.03654.',
        '[10] Hlavnova, E. and Pikuliak, M., "Fine-tuned LLMs for Multilingual Machine-Generated Text Detection," SemEval-2024 Task 8. arXiv:2402.13671.',
        '[11] "Fine-grained AI-generated Text Detection using Multi-task Auxiliary and Multi-level Contrastive Learning," (FAIDSet), 2025. arXiv:2505.14271.',
        '[12] Foltynek, T., Meuschke, N., and Gipp, B., "A Survey on Plagiarism Detection," ACM Computing Surveys, 2019. arXiv:1703.05546.',
        '[13] "Authorship Style Transfer with Policy Optimization," 2024. arXiv:2403.08043.',
    ]

    for ref in references:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(ref)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(4)

    # ─── Save ───
    output_path = 'Literature_Survey_Synopsis.docx'
    doc.save(output_path)
    print(f'Literature Survey (RV University Format) saved to: {output_path}')
    print(f'Contents:')
    print(f'  - Section 1: Paperwise comparison table (13 papers, 6 columns)')
    print(f'  - Section 2: Each paper elaborated in paragraph with arXiv hyperlinks')
    print(f'  - Section 3: References list')


if __name__ == '__main__':
    main()
