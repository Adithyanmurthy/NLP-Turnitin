"""
File Parser Module — Extracts text from any supported file format.
Supports: TXT, PDF, DOCX, DOC, MD, HTML, RTF

This is the first step in the pipeline — before any analysis happens,
the input file must be converted to clean plain text.
"""

import os
from pathlib import Path
from typing import Tuple


SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx", ".doc", ".md", ".html", ".htm", ".rtf"}


def detect_format(file_path: str) -> str:
    """Detect file format from extension."""
    ext = Path(file_path).suffix.lower()
    if ext in SUPPORTED_EXTENSIONS:
        return ext
    raise ValueError(
        f"Unsupported file format: '{ext}'. "
        f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
    )


def extract_from_txt(file_path: str) -> str:
    """Extract text from a plain text or markdown file."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def extract_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file using PyMuPDF (fitz)."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError(
            "PyMuPDF is required for PDF parsing. Install with: pip install PyMuPDF"
        )

    text_parts = []
    with fitz.open(file_path) as doc:
        for page_num, page in enumerate(doc):
            page_text = page.get_text("text")
            if page_text.strip():
                text_parts.append(page_text)

    return "\n\n".join(text_parts)


def extract_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX parsing. Install with: pip install python-docx"
        )

    doc = Document(file_path)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)

    return "\n\n".join(paragraphs)


def extract_from_html(file_path: str) -> str:
    """Extract text from an HTML file by stripping tags."""
    import re
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        html = f.read()

    # Remove script and style blocks
    html = re.sub(r"<script[^>]*>.*?</script>", "", html, flags=re.DOTALL | re.IGNORECASE)
    html = re.sub(r"<style[^>]*>.*?</style>", "", html, flags=re.DOTALL | re.IGNORECASE)
    # Remove tags
    text = re.sub(r"<[^>]+>", " ", html)
    # Clean whitespace
    text = re.sub(r"\s+", " ", text).strip()
    return text


def parse_file(file_path: str) -> Tuple[str, str]:
    """
    Parse any supported file and extract clean text.

    Args:
        file_path: Path to the input file.

    Returns:
        Tuple of (extracted_text, detected_format)

    Raises:
        FileNotFoundError: If file doesn't exist.
        ValueError: If format is unsupported.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    fmt = detect_format(file_path)

    if fmt in (".txt", ".md", ".rtf"):
        text = extract_from_txt(file_path)
    elif fmt == ".pdf":
        text = extract_from_pdf(file_path)
    elif fmt in (".docx", ".doc"):
        text = extract_from_docx(file_path)
    elif fmt in (".html", ".htm"):
        text = extract_from_html(file_path)
    else:
        raise ValueError(f"No parser available for format: {fmt}")

    # Basic cleaning
    text = text.strip()
    if not text:
        raise ValueError(f"No text content could be extracted from: {file_path}")

    return text, fmt


def parse_input(input_value: str) -> str:
    """
    Smart input parser — accepts either a file path or raw text.
    If the input looks like a file path and the file exists, parse it.
    Otherwise, treat it as raw text.

    Args:
        input_value: Either a file path or raw text string.

    Returns:
        Extracted/cleaned text string.
    """
    # Check if it's a file path
    if os.path.exists(input_value):
        try:
            text, fmt = parse_file(input_value)
            print(f"[FileParser] Parsed {fmt} file: {input_value} ({len(text)} chars)")
            return text
        except (ValueError, ImportError) as e:
            print(f"[FileParser] Could not parse file: {e}")
            print("[FileParser] Treating input as raw text")

    # It's raw text
    return input_value.strip()
